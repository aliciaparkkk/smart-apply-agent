from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
import copy
from pypdf import PdfReader

SUPPORTED_SUFFIXES = {".docx", ".pdf"}


def read_resume_lines(path: Path) -> list[str]:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        doc = Document(path)
        return [p.text for p in doc.paragraphs]
    if suffix == ".pdf":
        reader = PdfReader(path)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return text.splitlines()
    raise ValueError(f"Unsupported resume format: {suffix}. Use .docx or .pdf.")


def _set_paragraph_text_preserve_style(para, new_text: str) -> None:
    """Replace paragraph text while keeping the first run's formatting."""
    if not para.runs:
        para.text = new_text
        return

    # 첫 번째 run의 스타일 저장
    first_run = para.runs[0]
    bold = first_run.bold
    italic = first_run.italic
    underline = first_run.underline
    font_name = first_run.font.name
    font_size = first_run.font.size
    color = first_run.font.color.rgb if first_run.font.color and first_run.font.color.type else None

    # 모든 run 지우기
    for run in para.runs:
        run.text = ""

    # 첫 번째 run에만 새 텍스트 + 스타일 적용
    first_run.text = new_text
    first_run.bold = bold
    first_run.italic = italic
    first_run.underline = underline
    if font_name:
        first_run.font.name = font_name
    if font_size:
        first_run.font.size = font_size


def write_tailored_resume(
    source_path: Path,
    paragraphs: list[str],
    dest_path: Path,
) -> None:
    suffix = source_path.suffix.lower()
    if suffix == ".docx":
        doc = Document(source_path)
        if len(doc.paragraphs) != len(paragraphs):
            raise ValueError(
                f"Paragraph count mismatch: source has {len(doc.paragraphs)}, "
                f"tailored has {len(paragraphs)}"
            )
        for para, text in zip(doc.paragraphs, paragraphs):
            _set_paragraph_text_preserve_style(para, text)
        doc.save(dest_path)
        return

    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(dest_path)